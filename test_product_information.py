from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from reportlab.pdfgen import canvas

import product_information
from product_information import (
    facts_for_product,
    load_product_information_index,
    product_information_status,
    scan_product_information,
)


def _cfg(tmp_path: Path):
    return SimpleNamespace(
        WORKING_DIR=str(tmp_path / "working"),
        PRODUCT_INFORMATION_DIR=str(tmp_path / "assets" / "information"),
    )


def _write_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("PROYA TONER", level=1)
    document.add_heading("Ingredients", level=2)
    document.add_paragraph("Niacinamide, Panthenol")
    document.add_heading("Benefits", level=2)
    document.add_paragraph("Membantu menjaga kelembapan kulit")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cara pakai"
    table.cell(0, 1).text = "Gunakan setelah membersihkan wajah"
    document.add_heading("PROYA SERUM", level=1)
    document.add_paragraph("Ingredients: Vitamin C")
    document.save(path)


def _write_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 760, "PROYA CLEANSER")
    pdf.drawString(72, 730, "Ingredients: Amino Acid")
    pdf.drawString(72, 700, "Benefits: Membersihkan kulit dengan lembut")
    pdf.showPage()
    pdf.save()


def test_imports_docx_and_pdf_with_evidence(tmp_path):
    cfg = _cfg(tmp_path)
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    _write_docx(root / "catalog.docx")
    _write_pdf(root / "brochures" / "cleanser.pdf")

    index = scan_product_information(cfg)

    toner_facts = facts_for_product(index, "toner")
    cleanser_facts = facts_for_product(index, "cleanser")
    serum_facts = facts_for_product(index, "serum")
    assert {fact["role"] for fact in toner_facts} >= {"ingredient", "benefit", "usage"}
    assert any(fact["text"] == "Niacinamide" for fact in toner_facts)
    assert any(fact["locator"]["kind"] == "docx_table" for fact in toner_facts)
    assert any(fact["text"] == "Amino Acid" for fact in cleanser_facts)
    assert any(fact["locator"].get("page") == 1 for fact in cleanser_facts)
    assert any(fact["text"] == "Vitamin C" for fact in serum_facts)
    assert len(index["sources"]) == 2


def test_reuses_unchanged_sources_and_reports_bad_files(tmp_path):
    cfg = _cfg(tmp_path)
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    _write_docx(root / "toner.docx")
    (root / "broken.pdf").write_bytes(b"not a pdf")

    first = scan_product_information(cfg)
    second = load_product_information_index(cfg)
    status = product_information_status(cfg)

    assert first["revision"] == second["revision"]
    assert any(source["cached"] for source in second["sources"] if source["path"] == "toner.docx")
    assert any(source["status"] == "error" for source in status["sources"])


def test_empty_searchable_content_is_not_eligible(tmp_path):
    cfg = _cfg(tmp_path)
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    root.mkdir(parents=True)
    blank = canvas.Canvas(str(root / "blank.pdf"))
    blank.showPage()
    blank.save()

    status = product_information_status(cfg)

    assert status["sources"][0]["status"] == "image_only"
    assert status["products"][0]["eligible_fact_count"] == 0


def test_conflicting_product_identities_are_excluded(tmp_path):
    cfg = _cfg(tmp_path)
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    root.mkdir(parents=True)
    for filename, heading in (("a.docx", "PROYA TONER ALPHA"), ("b.docx", "PROYA TONER BETA")):
        document = Document()
        document.add_heading(heading, level=1)
        document.add_paragraph("Ingredients: Niacinamide")
        document.save(root / filename)

    index = scan_product_information(cfg)
    toner = index["products"]["toner"]

    assert any(item["reason"] == "product_identity_conflict" for item in index["conflicts"])
    assert not any(
        fact["eligible"]
        for fact in toner["facts"]
        if fact["role"] == "product_name"
    )


def test_llm_extracts_structured_fact_with_verbatim_evidence(tmp_path, monkeypatch):
    cfg = _cfg(tmp_path)
    cfg.PRODUCT_INFORMATION_LLM_ENABLED = True
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    root.mkdir(parents=True)
    document = Document()
    document.add_heading("PROYA TONER", level=1)
    document.add_paragraph("Niacinamide improves skin hydration")
    document.save(root / "catalog.docx")

    def fake_llm(payload, _cfg):
        assert payload["allowed_products"]
        return json.dumps({
            "facts": [{
                "block_id": 1,
                "product": "toner",
                "role": "benefit",
                "text": "Niacinamide improves skin hydration",
                "confidence": 0.96,
            }],
            "unassigned": [],
        })

    monkeypatch.setattr(product_information, "_call_information_llm", fake_llm)

    index = scan_product_information(cfg)
    facts = facts_for_product(index, "toner")

    assert len(facts) == 1
    assert facts[0]["extraction_method"] == "llm"
    assert facts[0]["source_excerpt"] == "Niacinamide improves skin hydration"
    assert index["sources"][0]["extraction_method"] == "llm"


def test_llm_rejects_fact_not_present_in_referenced_block(tmp_path, monkeypatch):
    cfg = _cfg(tmp_path)
    cfg.PRODUCT_INFORMATION_LLM_ENABLED = True
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    root.mkdir(parents=True)
    document = Document()
    document.add_heading("PROYA SERUM", level=1)
    document.add_paragraph("Contains Vitamin C")
    document.save(root / "serum.docx")

    monkeypatch.setattr(
        product_information,
        "_call_information_llm",
        lambda payload, _cfg: (
            '{"facts":[{"block_id":1,"product":"serum","role":"benefit",'
            '"text":"Erases wrinkles overnight","confidence":0.99}],"unassigned":[]}'
        ),
    )

    index = scan_product_information(cfg)

    assert facts_for_product(index, "serum") == []
    assert "rejected 1 unsupported or unverifiable item" in index["sources"][0]["warnings"][0]


def test_llm_failure_uses_rule_based_fallback(tmp_path, monkeypatch):
    cfg = _cfg(tmp_path)
    cfg.PRODUCT_INFORMATION_LLM_ENABLED = True
    root = Path(cfg.PRODUCT_INFORMATION_DIR)
    _write_docx(root / "toner.docx")

    def fail_llm(_payload, _cfg):
        raise TimeoutError("LM Studio unavailable")

    monkeypatch.setattr(product_information, "_call_information_llm", fail_llm)

    index = scan_product_information(cfg)
    status = product_information_status(cfg)

    assert facts_for_product(index, "toner")
    assert status["sources"][0]["extraction_method"] == "rules_fallback"
    assert any("rule-based fallback" in warning for warning in status["warnings"])
